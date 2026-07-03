import streamlit as st
import asyncio
import io
import os

# ── Trust the Windows certificate store (handles corporate SSL proxies) ────────
try:
    import truststore
    truststore.inject_into_ssl()
except Exception:
    pass  # Not on Windows or truststore not installed — continue without it

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Text Reader & Transcriber",
    page_icon="🎙️",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# ── Tesseract path detection (Windows local runs only) ─────────────────────────
TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(
        os.environ.get("USERNAME", "")
    ),
]

def find_tesseract():
    for path in TESSERACT_PATHS:
        if os.path.isfile(path):
            return path
    return None


# ── Extraction helpers ──────────────────────────────────────────────────────────

def extract_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a .docx file — covers:
      • Regular body paragraphs
      • Tables (rows × cells)
      • Text boxes / shapes (via fallback XML scan)
      • Headers and footers
    Uses the high-level python-docx API so it works on all Word file versions.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []

    # ── Walk body elements in document order (paragraphs + tables interleaved) ──
    for block in doc.element.body:
        # Plain paragraph
        if block.tag.endswith("}p"):
            text = block.text_content if hasattr(block, "text_content") else ""
            # Safer: join all w:t runs inside this paragraph
            text = "".join(t.text for t in block.iter() if t.tag.endswith("}t") and t.text)
            if text.strip():
                parts.append(text.strip())

        # Table
        elif block.tag.endswith("}tbl"):
            for row in block.iter():
                if not row.tag.endswith("}tr"):
                    continue
                cells = []
                for cell in row:
                    if not cell.tag.endswith("}tc"):
                        continue
                    cell_text = "".join(
                        t.text for t in cell.iter()
                        if t.tag.endswith("}t") and t.text
                    )
                    if cell_text.strip():
                        cells.append(cell_text.strip())
                if cells:
                    parts.append("  |  ".join(cells))

    # ── Headers and footers ───────────────────────────────────────────────────
    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                if hf is not None:
                    for para in hf.paragraphs:
                        if para.text.strip():
                            parts.append(para.text.strip())
            except Exception:
                pass  # missing header/footer is fine

    result = "\n".join(parts)

    # ── Last-resort fallback: if we got nothing, try doc.paragraphs directly ──
    if not result.strip():
        fallback = [p.text for p in doc.paragraphs if p.text.strip()]
        result = "\n".join(fallback)

    return result


def extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from every page of a PDF."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_from_image(file_bytes: bytes, lang: str = "eng") -> str:
    """Run Tesseract OCR on an image and return the recognised text."""
    import pytesseract
    from PIL import Image

    tess_path = find_tesseract()
    if tess_path:
        pytesseract.pytesseract.tesseract_cmd = tess_path
    else:
        st.error(
            "⚠️ Tesseract OCR not found. "
            "Please install it from https://github.com/UB-Mannheim/tesseract/wiki "
            "and restart the app."
        )
        return ""

    image = Image.open(io.BytesIO(file_bytes))
    return pytesseract.image_to_string(image, lang=lang)


# ── Edge TTS helpers ────────────────────────────────────────────────────────────

# Curated list of the best English neural voices available via edge-tts.
# Full name format used by the edge-tts library.
VOICES = {
    # ── English (US) ──
    "🇺🇸 Aria  — Female, warm & friendly":        "en-US-AriaNeural",
    "🇺🇸 Jenny — Female, conversational":          "en-US-JennyNeural",
    "🇺🇸 Michelle — Female, clear & professional": "en-US-MichelleNeural",
    "🇺🇸 Guy   — Male, casual":                    "en-US-GuyNeural",
    "🇺🇸 Davis — Male, energetic":                 "en-US-DavisNeural",
    "🇺🇸 Tony  — Male, calm":                      "en-US-TonyNeural",
    # ── English (UK) ──
    "🇬🇧 Libby  — Female, British":                "en-GB-LibbyNeural",
    "🇬🇧 Sonia  — Female, British professional":   "en-GB-SoniaNeural",
    "🇬🇧 Ryan   — Male, British":                  "en-GB-RyanNeural",
    # ── English (AU) ──
    "🇦🇺 Natasha — Female, Australian":            "en-AU-NatashaNeural",
    "🇦🇺 William — Male, Australian":              "en-AU-WilliamNeural",
}

def _fmt_rate(val: int) -> str:
    """edge-tts rate format: '+10%', '-5%', '+0%'"""
    return f"+{val}%" if val >= 0 else f"{val}%"

def _fmt_pitch(val: int) -> str:
    """edge-tts pitch format: '+10Hz', '-5Hz'  (NOT percent)"""
    return f"+{val}Hz" if val >= 0 else f"{val}Hz"


async def _edge_tts_generate(text: str, voice: str, rate: str, pitch: str) -> bytes:
    """Call edge-tts async API and return raw MP3 bytes."""
    import edge_tts
    # Only pass pitch when it's non-zero — edge-tts rejects '+0Hz' on some versions
    if pitch in ("+0Hz", "-0Hz", "0Hz"):
        communicate = edge_tts.Communicate(text, voice, rate=rate)
    else:
        communicate = edge_tts.Communicate(text, voice, rate=rate, pitch=pitch)
    buf = io.BytesIO()
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            buf.write(chunk["data"])
    return buf.getvalue()


def generate_audio(text: str, voice: str, rate_pct: int = 0, pitch_pct: int = 0) -> bytes:
    """Synchronous wrapper around the async edge-tts call."""
    rate  = _fmt_rate(rate_pct)
    pitch = _fmt_pitch(pitch_pct)
    return asyncio.run(_edge_tts_generate(text, voice, rate, pitch))


# ── Sidebar ─────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("🎙️ Text Reader")
    st.markdown(
        """
        **How to use:**
        1. Pick an input method (tabs below)
        2. Paste text or upload a file
        3. Click **Extract Text**
        4. Choose a voice, then click **🔊 Generate Audio**
        5. Press ▶ on the audio player
        """
    )
    st.divider()
    st.subheader("⚙️ Voice Settings")

    voice_label = st.selectbox(
        "Voice",
        options=list(VOICES.keys()),
        index=0,
    )
    selected_voice = VOICES[voice_label]

    rate_pct = st.slider(
        "Reading speed",
        min_value=-50, max_value=100, value=0, step=5,
        help="0% = normal speed. Negative = slower, positive = faster.",
        format="%d%%",
    )
    pitch_pct = st.slider(
        "Pitch",
        min_value=-50, max_value=50, value=0, step=5,
        help="0% = natural pitch. Negative = deeper, positive = higher.",
        format="%d%%",
    )

    st.divider()
    st.subheader("🔤 Other Settings")
    ocr_lang = st.selectbox(
        "OCR language (images)",
        options=["eng", "spa", "fra", "deu", "ita", "por", "chi_sim", "jpn"],
        format_func=lambda x: {
            "eng": "English", "spa": "Spanish", "fra": "French",
            "deu": "German",  "ita": "Italian", "por": "Portuguese",
            "chi_sim": "Chinese (Simplified)", "jpn": "Japanese",
        }.get(x, x),
    )
    st.divider()
    st.caption("Voices powered by Microsoft Edge Neural TTS — free, no API key needed.")


# ── Main area ───────────────────────────────────────────────────────────────────
st.title("🎙️ Text Reader & Transcriber")
st.caption("Paste text, upload a Word doc, PDF, or image — then have it read back to you.")

# ── Input tabs ──────────────────────────────────────────────────────────────────
tab_paste, tab_word, tab_pdf, tab_image = st.tabs(
    ["📋 Paste Text", "📄 Word Doc", "📑 PDF", "🖼️ Image / Photo"]
)

with tab_paste:
    st.subheader("Paste your text")
    st.caption("Copy from Google Docs, Word, email, websites, or anywhere else.")
    pasted = st.text_area(
        "Paste text here",
        height=250,
        placeholder="Paste any text here…",
        label_visibility="collapsed",
    )
    if st.button("Use This Text", key="use_paste", type="primary"):
        if pasted.strip():
            st.session_state["text"] = pasted.strip()
            st.session_state.pop("audio", None)
            st.success(f"✅ {len(pasted.split())} words loaded.")
        else:
            st.warning("Nothing to use — text area is empty.")

with tab_word:
    st.subheader("Upload a Word document (.docx)")
    st.caption("Must be a **.docx** file — not .doc (old format). Save as .docx from Word if needed.")
    docx_file = st.file_uploader("Choose a .docx file", type=["docx"], key="docx_up")
    if docx_file and st.button("Extract Text from Word Doc", key="ext_docx", type="primary"):
        with st.spinner("Reading document…"):
            try:
                result = extract_from_docx(docx_file.read())
                if not result.strip():
                    st.warning(
                        "The document was opened but no text was found. "
                        "Make sure the file is a standard .docx (not password-protected or corrupted)."
                    )
                else:
                    st.session_state["text"] = result
                    st.session_state.pop("audio", None)
                    st.success(f"✅ {len(result.split())} words extracted.")
            except Exception as e:
                import traceback
                st.error(
                    f"**Could not read file.**\n\n"
                    f"Error: `{e}`\n\n"
                    f"Full details:\n```\n{traceback.format_exc()}\n```"
                )

with tab_pdf:
    st.subheader("Upload a PDF")
    pdf_file = st.file_uploader("Choose a PDF file", type=["pdf"], key="pdf_up")
    if pdf_file and st.button("Extract Text from PDF", key="ext_pdf", type="primary"):
        with st.spinner("Reading PDF…"):
            try:
                result = extract_from_pdf(pdf_file.read())
                if not result.strip():
                    st.warning(
                        "No text found — this PDF may be image-based (scanned). "
                        "Try the Image tab instead."
                    )
                else:
                    st.session_state["text"] = result
                    st.session_state.pop("audio", None)
                    st.success(f"✅ {len(result.split())} words extracted.")
            except Exception as e:
                st.error(f"Could not read PDF: {e}")

with tab_image:
    st.subheader("Upload an image or photo")
    st.caption("Works with photos of documents, screenshots, printed text, signs, etc.")
    img_file = st.file_uploader(
        "Choose an image", type=["png", "jpg", "jpeg", "bmp", "tiff"], key="img_up"
    )
    if img_file:
        st.image(img_file, caption="Uploaded image", use_column_width=True)
    if img_file and st.button("Extract Text via OCR", key="ext_img", type="primary"):
        with st.spinner("Running OCR — this may take a few seconds…"):
            try:
                result = extract_from_image(img_file.read(), lang=ocr_lang)
                if not result.strip():
                    st.warning("No text detected in this image.")
                else:
                    st.session_state["text"] = result
                    st.session_state.pop("audio", None)
                    st.success(f"✅ {len(result.split())} words detected.")
            except Exception as e:
                st.error(f"OCR failed: {e}")


# ── Output area ─────────────────────────────────────────────────────────────────
st.divider()
st.subheader("📝 Extracted / Transcribed Text")

current_text = st.session_state.get("text", "")

if current_text:
    edited_text = st.text_area(
        "Text (editable — tweak before generating audio)",
        value=current_text,
        height=300,
        label_visibility="collapsed",
        key="text_display",
    )

    col1, col2, col3 = st.columns([2, 2, 1])
    with col1:
        word_count = len(edited_text.split())
        char_count = len(edited_text)
        st.caption(f"**{word_count:,} words · {char_count:,} characters**")
    with col3:
        if st.button("🗑️ Clear", key="clear_text"):
            st.session_state.pop("text", None)
            st.session_state.pop("audio", None)
            st.rerun()

    # Copy-to-clipboard
    safe_text = edited_text.replace("\\", "\\\\").replace("`", "\\`").replace("$", "\\$")
    copy_js = f"""
    <script>
      function copyText() {{
        navigator.clipboard.writeText(`{safe_text}`)
          .then(() => document.getElementById('copy-msg').textContent = '✅ Copied!')
          .catch(() => document.getElementById('copy-msg').textContent = '❌ Copy failed');
        setTimeout(() => document.getElementById('copy-msg').textContent = '', 2000);
      }}
    </script>
    <button onclick="copyText()"
      style="padding:8px 16px; font-size:14px; border-radius:6px;
             background:#1E88E5; color:white; border:none; cursor:pointer; margin-right:8px;">
      📋 Copy to Clipboard
    </button>
    <span id="copy-msg" style="font-size:13px; color:green;"></span>
    """
    st.components.v1.html(copy_js, height=50)

    # ── Audio generation ────────────────────────────────────────────────────────
    st.subheader("🔊 Read Aloud")
    st.caption(f"Voice: **{voice_label}** · Speed: **{_fmt_rate(rate_pct)}** · Pitch: **{_fmt_pitch(pitch_pct)}**")

    if word_count > 2000:
        st.warning(
            f"⏳ {word_count:,} words is a long read — audio generation may take "
            "30–60 seconds. You can also trim the text above and generate in sections."
        )

    if st.button("🔊 Generate Audio", type="primary", key="gen_audio"):
        with st.spinner(f"Generating audio with {voice_label.split('—')[0].strip()}…"):
            try:
                audio_bytes = generate_audio(
                    edited_text,
                    voice=selected_voice,
                    rate_pct=rate_pct,
                    pitch_pct=pitch_pct,
                )
                st.session_state["audio"] = audio_bytes
            except Exception as e:
                st.error(
                    f"Audio generation failed: {e}\n\n"
                    "Make sure you have an internet connection — "
                    "Microsoft Edge TTS requires it."
                )

    if "audio" in st.session_state:
        st.audio(st.session_state["audio"], format="audio/mp3")
        st.caption("▲ Tap ▶ to play. Works on all devices and mobile browsers.")
        st.download_button(
            label="⬇️ Download MP3",
            data=st.session_state["audio"],
            file_name="transcription.mp3",
            mime="audio/mpeg",
        )

else:
    st.info("👆 Choose an input method above, then click the Extract button to get started.")
